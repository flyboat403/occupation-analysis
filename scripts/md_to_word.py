#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word converter (simple version)
"""

import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_word_report(data_path: str, output_path: str):
    """从JSON数据生成Word报告"""
    # 加载数据
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 创建文档
    doc = Document()
    
    # 标题
    title = doc.add_heading(data['meta']['report_title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph(f"专业名称：{data['meta']['major_name']}")
    doc.add_paragraph(f"专业代码：{data['meta']['major_code']}")
    doc.add_paragraph(f"教育层次：{data['meta']['education_level']}")
    doc.add_paragraph(f"生成日期：{data['meta']['generated_date']}")
    doc.add_paragraph()
    
    # 第一部分：职业面向
    doc.add_heading('一、职业面向', level=1)
    doc.add_paragraph(f"专业名称：{data['occupation_facing']['major_name']}")
    doc.add_paragraph(f"专业代码：{data['occupation_facing']['major_code']}")
    
    # 表1：职业面向表
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Light Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '岗位名称'
    hdr_cells[1].text = '中国职业代码'
    hdr_cells[2].text = 'ISCO代码'
    hdr_cells[3].text = 'O*NET代码'
    
    for pos in data['occupation_facing']['positions']:
        row_cells = table1.add_row().cells
        row_cells[0].text = pos['position_name']
        row_cells[1].text = pos['occupation_code']
        row_cells[2].text = pos['isco_code']
        row_cells[3].text = pos['onet_code']
    
    doc.add_paragraph()
    
    # 第二部分：典型工作任务
    doc.add_heading('二、典型工作任务分析', level=1)
    
    # 表2：典型工作任务汇总表
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Light Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '典型工作任务'
    hdr_cells[2].text = '难度等级'
    hdr_cells[3].text = '关联岗位'
    hdr_cells[4].text = '工作对象'
    
    for i, task in enumerate(data['typical_tasks'], 1):
        row_cells = table2.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = task['name']
        row_cells[2].text = task['difficulty']
        row_cells[3].text = task['position']
        row_cells[4].text = task['work_object']
    
    doc.add_paragraph()
    
    # 第三部分：行动领域
    doc.add_heading('三、行动领域划分', level=1)
    
    # 表4：行动领域表
    table4 = doc.add_table(rows=1, cols=4)
    table4.style = 'Light Grid'
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '行动领域'
    hdr_cells[2].text = '包含的典型任务'
    hdr_cells[3].text = '能力等级'
    
    for i, domain in enumerate(data['action_domains'], 1):
        row_cells = table4.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = domain['name']
        row_cells[2].text = ', '.join(domain['tasks'])
        row_cells[3].text = '初级' if '维护' in domain['name'] else ('高级' if '新能源' in domain['name'] else '中级')
    
    doc.add_paragraph()
    
    # 第四部分：职业能力分析
    doc.add_heading('四、职业能力分析', level=1)
    
    # 表6：职业能力分析汇总表
    table6 = doc.add_table(rows=1, cols=4)
    table6.style = 'Light Grid'
    hdr_cells = table6.rows[0].cells
    hdr_cells[0].text = '行动领域'
    hdr_cells[1].text = '专业能力'
    hdr_cells[2].text = '方法能力'
    hdr_cells[3].text = '社会能力'
    
    for domain in data['action_domains']:
        row_cells = table6.add_row().cells
        row_cells[0].text = domain['name']
        row_cells[1].text = domain['abilities']['professional']
        row_cells[2].text = domain['abilities']['method']
        row_cells[3].text = domain['abilities']['social']
    
    doc.add_paragraph()
    
    # 第五部分：学习领域
    doc.add_heading('五、学习领域设计', level=1)
    
    # 表9：学习领域表
    table9 = doc.add_table(rows=1, cols=5)
    table9.style = 'Light Grid'
    hdr_cells = table9.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '学习领域'
    hdr_cells[2].text = '参考学时'
    hdr_cells[3].text = '学习目标'
    hdr_cells[4].text = '学习内容'
    
    total_hours = 0
    for i, field in enumerate(data['learning_domains'], 1):
        row_cells = table9.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = field['name']
        row_cells[2].text = str(field['hours'])
        row_cells[3].text = field['objective']
        row_cells[4].text = field['content']
        total_hours += field['hours']
    
    doc.add_paragraph()
    doc.add_paragraph(f"总计学时：{total_hours} 学时", style='Intense Quote')
    doc.add_paragraph()
    
    # 第六部分：学习情境
    doc.add_heading('六、学习情境设计', level=1)
    
    for field in data['learning_domains']:
        doc.add_heading(field['name'], level=2)
        doc.add_paragraph(f"参考学时：{field['hours']}学时")
        doc.add_paragraph(f"学习目标：{field['objective']}")
        
        # 表10：学习情境表
        table10 = doc.add_table(rows=1, cols=4)
        table10.style = 'Light Grid'
        hdr_cells = table10.rows[0].cells
        hdr_cells[0].text = '学习情境'
        hdr_cells[1].text = '学时'
        hdr_cells[2].text = '教学方法'
        hdr_cells[3].text = '评价方式'
        
        for situation in field.get('situations', []):
            row_cells = table10.add_row().cells
            row_cells[0].text = situation['name']
            row_cells[1].text = str(situation['hours'])
            row_cells[2].text = situation['methods']
            row_cells[3].text = situation['assessment']
        
        doc.add_paragraph()
    
    # 第七部分：国际职业信息
    doc.add_heading('七、国际职业信息对照', level=1)
    
    int_data = data.get('international_data', {}).get('occupation_info', {})
    china = int_data.get('china_occupation', {})
    isco = int_data.get('isco_occupation', {})
    onet = int_data.get('onet_occupation', {})
    
    doc.add_paragraph('中国职业大典：', style='Heading 2')
    doc.add_paragraph(f"职业名称：{china.get('name', '')}")
    doc.add_paragraph(f"职业代码：{china.get('code', '')}")
    doc.add_paragraph(f"职业描述：{china.get('description', '')}")
    
    doc.add_paragraph('ESCO职业分类：', style='Heading 2')
    doc.add_paragraph(f"职业名称：{isco.get('name', '')} ({isco.get('name_zh', '')})")
    doc.add_paragraph(f"职业代码：{isco.get('code', '')}")
    
    doc.add_paragraph('O*NET职业分类：', style='Heading 2')
    doc.add_paragraph(f"职业名称：{onet.get('name', '')} ({onet.get('name_zh', '')})")
    doc.add_paragraph(f"职业代码：{onet.get('code', '')}")
    
    # 保存文档
    doc.save(output_path)
    print(f"[OK] Word报告已生成：{output_path}")


if __name__ == '__main__':
    create_word_report(
        data_path='analysis_data.json',
        output_path='output/report.docx'
    )
